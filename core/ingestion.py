"""
Document ingestion — extract text from PDF, DOCX, and TXT files,
then split into overlapping chunks for embedding.
"""
from __future__ import annotations

import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Union

from langchain_text_splitters import RecursiveCharacterTextSplitter

# ── Constants ────────────────────────────────────────────────────────────────
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
SEPARATORS = ["\n\n", "\n", ". ", " ", ""]

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}


# ── Extractors ────────────────────────────────────────────────────────────────
def extract_text_from_pdf(source: Union[BytesIO, Path, str]) -> str:
    """Extract all text from a PDF file or file-like object."""
    from pypdf import PdfReader
    reader = PdfReader(source)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(source: Union[BytesIO, bytes, Path, str]) -> str:
    """Extract paragraph text from a DOCX file or bytes."""
    from docx import Document

    if isinstance(source, (bytes, BytesIO)):
        data = source if isinstance(source, bytes) else source.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        doc = Document(tmp_path)
        os.unlink(tmp_path)
    else:
        doc = Document(str(source))

    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_txt(source: Union[BytesIO, bytes, Path, str]) -> str:
    """Decode plain text from a file or bytes."""
    if isinstance(source, (bytes, BytesIO)):
        data = source if isinstance(source, bytes) else source.read()
        return data.decode("utf-8", errors="replace")
    return Path(str(source)).read_text(encoding="utf-8", errors="replace")


_EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt,
}


def extract_text(file, mime_type: str | None = None, filename: str | None = None) -> str:
    """
    Dispatch to the correct extractor.

    Parameters
    ----------
    file        : file-like object, bytes, or path
    mime_type   : MIME type string (e.g. "application/pdf")
    filename    : fallback for extension detection when mime_type is absent
    """
    ext: str | None = None

    if mime_type:
        ext = MIME_TYPES.get(mime_type)

    if ext is None and filename:
        ext = Path(filename).suffix.lower()

    extractor = _EXTRACTORS.get(ext or "")
    if extractor is None:
        raise ValueError(
            f"Unsupported file type '{ext or mime_type}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    return extractor(file)


# ── Chunking ─────────────────────────────────────────────────────────────────
def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    """Split *text* into overlapping chunks suitable for embedding."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=SEPARATORS,
    )
    return splitter.split_text(text)


# ── Combined pipeline ─────────────────────────────────────────────────────────
def load_and_chunk(files) -> tuple[list[str], list[str]]:
    """
    Extract text from each uploaded Streamlit file, combine, and chunk.

    Parameters
    ----------
    files : list of Streamlit UploadedFile objects

    Returns
    -------
    chunks      : list of text chunks ready for embedding
    file_names  : names of files that were successfully processed
    """
    all_texts: list[str] = []
    file_names: list[str] = []

    for f in files:
        try:
            text = extract_text(f, mime_type=getattr(f, "type", None), filename=getattr(f, "name", None))
            if text.strip():
                all_texts.append(text)
                file_names.append(f.name)
        except ValueError:
            pass  # Skip unsupported files silently

    if not all_texts:
        return [], []

    combined = "\n\n".join(all_texts)
    chunks = chunk_text(combined)
    return chunks, file_names
